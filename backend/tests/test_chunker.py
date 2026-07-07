"""
Tests for the chapter chunker service.
No IBM credentials required.
"""

import pytest

from app.services.chunker import ChapterChunk, parse_text, split_into_chapters

# ---------------------------------------------------------------------------
# parse_text
# ---------------------------------------------------------------------------

def test_parse_text_txt():
    content = "Hello world"
    result = parse_text(content.encode("utf-8"), "sample.txt")
    assert result == "Hello world"


def test_parse_text_txt_latin1_fallback():
    content = "Caf\xe9 au lait"  # latin-1 encoded
    result = parse_text(content.encode("latin-1"), "sample.txt")
    assert "Caf" in result


def test_parse_text_docx():
    """Create a minimal .docx in memory and verify extraction."""
    from io import BytesIO
    from docx import Document

    doc = Document()
    doc.add_paragraph("Chapter 1: The Beginning")
    doc.add_paragraph("It was a dark and stormy night.")
    buf = BytesIO()
    doc.save(buf)
    result = parse_text(buf.getvalue(), "sample.docx")
    assert "dark and stormy" in result


# ---------------------------------------------------------------------------
# split_into_chapters — heading detection
# ---------------------------------------------------------------------------

_HEADED_TEXT = """\
Chapter 1: The Beginning
It was a dark and stormy night.
The wind howled through the trees.

Chapter 2: Rising Action
Things got worse after that.
Much worse, in fact.

Chapter 3: The Climax
Everything came to a head.
"""


def test_heading_detection_returns_three_chapters():
    chunks = split_into_chapters(_HEADED_TEXT)
    assert len(chunks) == 3


def test_chapter_ids_are_stable():
    chunks = split_into_chapters(_HEADED_TEXT)
    assert chunks[0].chapter_id == "chapter_01"
    assert chunks[1].chapter_id == "chapter_02"
    assert chunks[2].chapter_id == "chapter_03"


def test_chapter_titles_captured():
    chunks = split_into_chapters(_HEADED_TEXT)
    assert "Beginning" in chunks[0].title


def test_word_counts_positive():
    chunks = split_into_chapters(_HEADED_TEXT)
    for chunk in chunks:
        assert chunk.word_count > 0


def test_text_included_in_chunk():
    chunks = split_into_chapters(_HEADED_TEXT)
    assert "dark and stormy" in chunks[0].text


# ---------------------------------------------------------------------------
# split_into_chapters — fallback (word count)
# ---------------------------------------------------------------------------

def test_fallback_splits_long_text():
    # 3001 words → should produce at least 2 chunks
    text = " ".join([f"word{i}" for i in range(3001)])
    chunks = split_into_chapters(text)
    assert len(chunks) >= 2


def test_fallback_chunk_ids():
    text = " ".join([f"word{i}" for i in range(1600)])
    chunks = split_into_chapters(text)
    assert chunks[0].chapter_id == "chapter_01"


def test_fallback_scene_titles():
    text = " ".join([f"word{i}" for i in range(1600)])
    chunks = split_into_chapters(text)
    assert chunks[0].title.startswith("Scene")
